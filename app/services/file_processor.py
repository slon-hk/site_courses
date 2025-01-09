import os
from docx import Document
import PyPDF2
import logging
from app.services.vector_db import VectorDB

logger = logging.getLogger(__name__)

class FileProcessor:
    _vector_db = None

    @classmethod
    def get_vector_db(cls):
        if cls._vector_db is None:
            cls._vector_db = VectorDB()
        return cls._vector_db

    @staticmethod
    def process_file(file_path):
        """Извлекает текст из файла и создает векторное представление"""
        try:
            # Получаем расширение файла
            file_ext = os.path.splitext(file_path)[1].lower()

            # Извлекаем текст в зависимости от типа файла
            text = ""
            try:
                if file_ext == '.docx':
                    text = FileProcessor._process_docx(file_path)
                elif file_ext == '.pdf':
                    text = FileProcessor._process_pdf(file_path)
                else:
                    raise ValueError(f"Неподдерживаемый тип файла: {file_ext}")
            except Exception as e:
                logger.error(f"Ошибка при извлечении текста из файла {file_path}: {str(e)}")
                text = "Ошибка при извлечении текста"

            # Создаем векторное представление
            try:
                vector_db = FileProcessor.get_vector_db()
                vector = vector_db.create_embedding(text)

                # Добавляем документ в векторную базу
                vector_db.add_document(text, file_path)

                return vector
            except Exception as e:
                logger.error(f"Ошибка при создании векторного представления: {str(e)}")
                return []

        except Exception as e:
            logger.error(f"Ошибка при обработке файла {file_path}: {str(e)}")
            return []

    @staticmethod
    def _process_docx(file_path):
        """Извлекает текст из DOCX файла"""
        try:
            text_parts = []
            doc = Document(file_path)

            # Извлекаем текст из параграфов
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Извлекаем текст из таблиц
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)

            text = '\n'.join(text_parts)
            if not text.strip():
                return "Документ пуст или не содержит текста"
            return text
        except Exception as e:
            logger.error(f"Ошибка при обработке DOCX файла: {str(e)}")
            return "Ошибка при чтении DOCX файла"

    @staticmethod
    def _process_pdf(file_path):
        """Извлекает текст из PDF файла"""
        try:
            text_parts = []
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text = page.extract_text()
                    if text.strip():
                        text_parts.append(text)

            text = '\n'.join(text_parts)
            if not text.strip():
                return "PDF документ пуст или не содержит текста"
            return text
        except Exception as e:
            logger.error(f"Ошибка при обработке PDF файла: {str(e)}")
            return "Ошибка при чтении PDF файла"

    @staticmethod
    def search_similar_documents(query, top_k=3):
        """Поиск похожих документов"""
        try:
            vector_db = FileProcessor.get_vector_db()
            return vector_db.search(query, top_k)
        except Exception as e:
            logger.error(f"Ошибка при поиске документов: {str(e)}")
            return []